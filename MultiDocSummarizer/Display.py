from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
import docx
from SummaryGenerator import SummaryGenerator

import tkinter as tk

class MyDialog:
    def __init__(self, parent,total_num_of_sent):
        top = self.top = tk.Toplevel(parent)
        top.configure(width=500,height = 300)
        self.myLabel = tk.Label(top, text='Lines in Summary')
        self.myLabel.place(relheight=0.1,x=5,y=10,relwidth=0.98)

        self.slide = Scale(top, orient=HORIZONTAL)
        self.slide.config(from_=0, to=total_num_of_sent)
        self.slide.place(relheight=0.2, x=5, y=30, relwidth=0.98)

        self.keys = tk.Text(top,wrap='word', font='Consolas 11',pady=0)
        self.keys.place(relheight=0.5, x=5, y=80, relwidth=0.96)
        self.keys.insert(1.0, 'Enter Keywords')

        self.mySubmitButton = tk.Button(top, text='Submit', command=self.send)
        self.mySubmitButton.place(x=220, y=260)
    def send(self):
        self.length = self.slide.get()
        self.keywords = self.keys.get(1.0,END)
        print("blbblb",self.length,self.keywords)
        self.top.destroy()

class Display:
    document_original = []

    ftypes = [
        ('Document Files', '*.docx;*.doc'),# semicolon trick
        ('Text Files', '*.txt'),
        ('Pdf Files', '*.pdf'),
        ('All files', '*')]

    window = Tk()

    def setup(self):

        self.window.state("zoomed")
        window_height = self.window.winfo_screenheight()-80#.winfo_height()
        window_width = self.window.winfo_screenwidth()#.winfo_geometry()
        self.window.wm_minsize(width=window_width//2-10,height=window_height-100)
        print(window_width)
        document_labelframe = LabelFrame(self.window,text="Document Text",bg="white")
        left_frame = Frame(self.window,bg="gray",width=window_width//2,padx=10)
        right_frame = Frame(self.window,bg="blue",width=window_width//2,padx=10,pady=10)


        text="Select a Document to view text"


        self.w = Text( right_frame,wrap='word', font='Consolas 11',pady=0)

        scrollbar2 = Scrollbar(self.w)
        scrollbar2.pack(side=RIGHT,fill=Y)
        scrollbar2.config(command=self.w.yview)
        self.w.config(yscrollcommand = scrollbar2.set)

        self.w.insert( 1.0, text )
        self.w.config(state=DISABLED)
        self.w.place(relwidth=1,relheight=1,relx=0,rely=0)


        self.document_list = Listbox(left_frame,width=window_width//2-40)

        scrollbar = Scrollbar(self.document_list)
        scrollbar.pack(side=RIGHT,fill=Y)
        scrollbar.config(command=self.document_list.yview)
        self.document_list.config(yscrollcommand = scrollbar.set)

        self.document_list.place(relheight=0.7,x=5,y=10,relwidth=0.98)

        bn_add_document = Button(left_frame, text="Add Document", command=self.addDocument)

        bn_add_document.place(relx=0.1, rely=0.75)
        bn_delete_document = Button(left_frame, text="Remove Document", command=self.deleteDocument)
        bn_delete_document.place(relx=0.7, rely=0.75)
        bn_generate_summary = Button(left_frame, text="Generate Summary", command=self.generateSummary)
        bn_generate_summary.place(relx=0.4, rely=0.8)

        right_frame.place(relx=0.4, relwidth=0.6, relheight=1)
        left_frame.place(relx=0, relwidth=0.4, relheight=1)

        self.document_list.bind('<<ListboxSelect>>', self.onselect)
        self.window.mainloop()

    def displayDocText(self,doc_text):
        self.w.config(state=NORMAL)
        self.w.delete(1.0, END)
        self.w.insert(1.0, doc_text)
        self.w.config(state=DISABLED)

    def addDocument(self):
        filenames = filedialog.askopenfilenames(filetypes=self.ftypes)
        if(filenames):
            for filename in filenames:
                self.document_list.insert(END,filename)
                doc = docx.Document(filename)
                doc_text = "\n".join([para.text for para in doc.paragraphs if (para.text).strip()!=''])
                self.document_original.append(doc_text)

    def deleteDocument(self):
        selection = self.document_list.curselection()
        if(not selection):
            messagebox.showwarning("Remove Document","Select a document to remove")
            print("Select an document to remove")
        else:
            if messagebox.askyesno("Remove Document", "Are you sure you want to remove this document"):
                # are you sure you want to remove this document
                doc = self.document_list.get(ACTIVE)
                self.document_list.delete(selection[0])
                # if document which is deleted is the one whose text is displayed then clear the texts
                print("remove",doc)
                self.displayDocText("")

    def generateSummary(self):
        print("Summarize : ")
        documents = self.document_list.get(0,END)
        summaryGenerator = SummaryGenerator()
        summaryGenerator.setDocPaths(documents)

        inputDialog = MyDialog(self.window,summaryGenerator.total_num_of_sent)
        self.window.wait_window(inputDialog.top)

        summaryGenerator.lines_in_summary = inputDialog.length
        summaryGenerator.title = inputDialog.keywords
        print(summaryGenerator.total_num_of_sent)

        summary = summaryGenerator.summarize()
        self.displayDocText(summary)

    def onselect(self,evt):
        selection = self.document_list.curselection()
        if(selection):
            print("Selected : ",selection[0])
            self.displayDocText(self.document_original[selection[0]])


